import os
from typing import Generator, List, Tuple, Dict, Any
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from models import PageLayout, TextBlock, WordBox
from logger import get_logger

logger = get_logger("parser")

class DocxNormalizer:
    """
    Unified Ingestion Layer for Microsoft Word (.docx) documents.
    Normalizes paragraphs and tables into standard PageLayout and TextBlock objects.
    """
    def __init__(self, config=None) -> None:
        self.config = config

    def _iter_block_items(self, parent: DocxDocument):
        """
        Generate each paragraph or table in the document body in order.
        """
        if isinstance(parent, DocxDocument):
            parent_elm = parent.element.body
        elif hasattr(parent, '_element'):
            parent_elm = parent._element
        else:
            return

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    def _extract_table_text(self, table: Table) -> List[str]:
        """
        Extracts all cell text sequences from a table.
        For a single cell, returns paragraphs inside.
        For multi-column tables, returns serialized rows.
        """
        lines = []
        # If it's a single cell table (layout box)
        if len(table.rows) == 1 and len(table.columns) == 1:
            cell = table.rows[0].cells[0]
            for p in cell.paragraphs:
                t = p.text.strip()
                if t:
                    lines.append(t)
        else:
            # Data table layout (e.g. data tables/DI)
            for r_idx, row in enumerate(table.rows):
                row_cells_text = []
                for cell in row.cells:
                    # Collect paragraph texts
                    cell_p_text = " ".join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                    row_cells_text.append(cell_p_text)
                # Serialize row as tab/pipe separated values
                if any(row_cells_text):
                    lines.append(" | ".join(row_cells_text))
        return lines

    def process_document(self, docx_path: str) -> Generator[PageLayout, None, None]:
        """
        Processes a DOCX file and yields PageLayout blocks sequentially.
        Simulates page pagination every 25 layout blocks or on explicit page breaks.
        """
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"DOCX file not found: {docx_path}")

        logger.info(f"Normalizing DOCX document: {docx_path}")
        doc = Document(docx_path)
        
        blocks_in_page: List[TextBlock] = []
        current_page_num = 1
        y_cursor = 50.0

        def create_word_boxes(text: str, page_num: int, start_y: float) -> List[WordBox]:
            words = text.split()
            boxes = []
            x_cursor = 50.0
            for idx, w in enumerate(words):
                w_len = len(w) * 6.0
                boxes.append(WordBox(
                    text=w,
                    x0=x_cursor,
                    y0=start_y,
                    x1=x_cursor + w_len,
                    y1=start_y + 12.0,
                    font_name="Calibri",
                    font_size=11.0,
                    page=page_num
                ))
                x_cursor += w_len + 5.0
            return boxes

        def create_text_block(text: str, page_num: int, start_y: float) -> TextBlock:
            words = create_word_boxes(text, page_num, start_y)
            return TextBlock(
                text=text,
                words=words,
                bbox=(50.0, start_y, 550.0, start_y + 15.0),
                page=page_num,
                block_type="text"
            )

        def yield_page() -> PageLayout:
            nonlocal blocks_in_page, current_page_num
            layout = PageLayout(
                page_number=current_page_num,
                width=595.0,
                height=842.0,
                rotation=0,
                layout_type="single_column",
                blocks=list(blocks_in_page),
                metadata={"engine": "python-docx"}
            )
            blocks_in_page.clear()
            current_page_num += 1
            return layout

        # Sequential traversal
        for item in self._iter_block_items(doc):
            lines_to_add = []
            if isinstance(item, Paragraph):
                text = item.text.strip()
                if text:
                    lines_to_add.append(text)
            elif isinstance(item, Table):
                table_lines = self._extract_table_text(item)
                lines_to_add.extend(table_lines)

            for line in lines_to_add:
                # Detect manual page breaks (e.g. section headers or explicit Word breaks)
                is_new_section = any(h in line.lower() for h in ["english", "quantitative", "reasoning", "general awareness", "computer"])
                
                # Yield current page if limit exceeded or new section starts
                if (len(blocks_in_page) >= 25 or is_new_section) and blocks_in_page:
                    yield yield_page()
                    y_cursor = 50.0

                block = create_text_block(line, current_page_num, y_cursor)
                blocks_in_page.append(block)
                y_cursor += 30.0

        if blocks_in_page:
            yield yield_page()

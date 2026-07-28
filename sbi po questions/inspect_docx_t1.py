import docx
import re

doc_path = r"C:\Users\Administrator\Downloads\sbi po prelims\SBI-PO-Pre-2022-19th-Dec-Shift-Wise-Previous-Year-Paper-Mock-5.docx"
doc = docx.Document(doc_path)

print("--- Paragraphs ---")
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if re.search(r'\bQ(10|11|12|13|35|36|37|45|46|50|51)\b', txt, re.IGNORECASE):
        print(f"P{i}: {txt[:120]}")

print("\n--- Tables ---")
for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            txt = cell.text.strip()
            if re.search(r'\bQ(10|11|12|13|35|36|37|45|46|50|51)\b', txt, re.IGNORECASE):
                print(f"Table {t_idx} R{r_idx} C{c_idx}: {txt[:120]}")


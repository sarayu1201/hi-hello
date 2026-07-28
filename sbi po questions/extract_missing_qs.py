import docx
import re
import json

docx_dir = r"C:\Users\Administrator\Downloads\sbi po prelims"

missing_specs = {
    1: [10, 11, 12, 13],
    2: [10, 12, 13, 14],
    4: [1, 2, 3, 4, 5, 57, 58],
    7: [22, 23, 24, 25],
    8: [1, 2, 3, 4, 61]
}

docx_files = [
    "SBI-PO-Pre-2022-19th-Dec-Shift-Wise-Previous-Year-Paper-Mock-5.docx",
    "SBI-PO-Pre-2022-20th-Dec-Shift-Wise-Previous-Year-Paper-Mock-6.docx",
    "SBI-PO-Pre-2024-25-Memory-Based-Paper-16-Mar-2025-1st-shift.docx",
    "SBI-PO-Pre-2024-25-Memory-Based-Paper-24-Mar-2025-1st-shift-1.docx",
    "SBI-PO-Pre-2024-25-Memory-Based-Paper-8-Mar-2025-1st-shift.docx",
    "SBI-PO-Pre-2024-25-Memory-Based-Paper-8-March-2025-3rd-shift.docx",
    "SBI-PO-Pre-2024-25-Memory-Based-Paper-8-March-2025-4th-shift.docx",
    "SBI-PO-Pre-2024-25-Memory-Based-Question-Paper-8-Mar-2025-2nd-shift-1.docx",
    "SBI-PO-Pre-2025-Memory-Based-Paper-Based-on-4th-August-1st-Shift (1).docx",
    "SBI-PO-Pre-2025-Memory-Based-Paper-Based-on-4th-August-1st-Shift.docx"
]

for test_num, q_ids in missing_specs.items():
    docx_name = docx_files[test_num - 1]
    doc_path = f"{docx_dir}\\{docx_name}"
    doc = docx.Document(doc_path)
    
    print(f"\n================ Test {test_num} Missing Qs: {q_ids} ================")
    
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    for qid in q_ids:
        # Search for Q{qid} in paragraphs
        found = False
        pattern = rf'^\s*Q\s*\.?\s*{qid}\b|^\s*{qid}\s*\.'
        for idx, p in enumerate(paragraphs):
            if re.search(pattern, p, re.IGNORECASE):
                print(f"\n--- Found Q{qid} in Paragraph {idx} ---")
                print("\n".join(paragraphs[max(0, idx-1):min(len(paragraphs), idx+6)]))
                found = True
                break
        if not found:
            # Check tables
            for t_idx, table in enumerate(doc.tables):
                for r_idx, row in enumerate(table.rows):
                    for c_idx, cell in enumerate(row.cells):
                        if re.search(pattern, cell.text, re.IGNORECASE):
                            print(f"\n--- Found Q{qid} in Table {t_idx} R{r_idx} C{c_idx} ---")
                            print(cell.text[:300])
                            found = True
                            break
                    if found: break
                if found: break


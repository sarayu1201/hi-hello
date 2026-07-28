import docx
import re

def parse_docx_questions(docx_path):
    doc = docx.Document(docx_path)
    text_blocks = []
    
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            text_blocks.append(txt)
            
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                txt = c.text.strip()
                if txt:
                    text_blocks.append(txt)
                    
    full_text = "\n\n".join(text_blocks)
    
    # Split by Q\d+
    q_dict = {}
    matches = list(re.finditer(r'\bQ(\d+)[\.\s]*', full_text, re.IGNORECASE))
    
    for idx, m in enumerate(matches):
        q_num = int(m.group(1))
        start_pos = m.end()
        end_pos = matches[idx+1].start() if idx + 1 < len(matches) else len(full_text)
        
        q_chunk = full_text[start_pos:end_pos].strip()
        
        # extract options
        opt_matches = list(re.finditer(r'\(([a-eA-E])\)\s*', q_chunk))
        if opt_matches:
            q_text = q_chunk[:opt_matches[0].start()].strip()
            options = {}
            for o_idx, om in enumerate(opt_matches):
                o_char = om.group(1).upper()
                o_start = om.end()
                o_end = opt_matches[o_idx+1].start() if o_idx + 1 < len(opt_matches) else len(q_chunk)
                o_val = q_chunk[o_start:o_end].strip()
                options[o_char] = o_val
            q_dict[q_num] = {
                "question": q_text,
                "options": options
            }
        else:
            q_dict[q_num] = {
                "question": q_chunk,
                "options": {}
            }
            
    return q_dict

# Test on test 1
res = parse_docx_questions(r"C:\Users\Administrator\Downloads\sbi po prelims\SBI-PO-Pre-2022-19th-Dec-Shift-Wise-Previous-Year-Paper-Mock-5.docx")
print(f"Parsed {len(res)} questions from DOCX Test 1.")
if 35 in res:
    print("Q35 from DOCX:", res[35])
if 45 in res:
    print("Q45 from DOCX:", res[45])

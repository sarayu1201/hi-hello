import json
import os
import re
import docx
import sys

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

orig_dir = r"C:\Users\Administrator\Downloads\question papers\json_output\sbi_po_prelims"
target_dir = r"C:\Users\Administrator\Downloads\hi-hello-main (1)\hi-hello-main\QuestionBank\json\sbi_po_prelims"
docx_dir = r"C:\Users\Administrator\Downloads\sbi po prelims"

docx_files = [
    "SBI-PO-Pre-2022-19th-Dec-Shift-Wise-Previous-Year-Paper-Mock-5.docx",
    "SBI-PO-Pre-2022-20th-Dec-Shift-Wise-Previous-Year-Paper-Mock-6.docx",
    "SBI-PO-Pre-2024-25-Memory-Based-Paper-16-Mar-2025-1st-shift.docx",
    "SBI-PO-Pre-2024-25-Memory-Based-Paper-24-Mar-2025-1st-shift-1.docx",
    "SBI-PO-Pre-2024-25-Memory-Based-Paper-8-Mar-2025-1st-shift.docx",
    "SBI-PO-Pre-2024-25-Memory-Based-Paper-8-March-2025-3rd-shift.docx",
    "SBI-PO-Pre-2024-25-Memory-Based-Paper-8-March-2025-4th-shift.docx",
    "SBI-PO-Pre-2024-25-Memory-Based-Question-Paper-8-Mar-2025-2nd-shift-1.docx",
    "SBI-PO-Pre-2025-Memory-Based-Paper-Based-on-4th-August-1st-Shift (1).docx",
    "SBI-PO-Pre-2025-Memory-Based-Paper-Based-on-4th-August-1st-Shift.docx"
]

DIR_OPTIONS_59_61 = {
    'A': '$x > y$',
    'B': '$x \\ge y$',
    'C': '$x < y$',
    'D': '$x \\le y$',
    'E': '$x = y$ or no relation can be established between $x$ and $y$'
}

DIR_OPTIONS_SYLLOGISM = {
    'A': 'If only conclusion I follows',
    'B': 'If only conclusion II follows',
    'C': 'If either conclusion I or II follows',
    'D': 'If neither conclusion I nor II follows',
    'E': 'If both conclusions I and II follow'
}

TEST2_MANUAL_OPTIONS = {
    24: {'A': 'Only (A)-(C)', 'B': 'Only (B)-(C) and (A)-(D)', 'C': 'Only (A)-(B) and (C)-(D)', 'D': 'Only (A)-(D)', 'E': 'None of these'},
    25: {'A': 'Only (A)-(D)', 'B': 'Only (B)-(A) and (C)-(D)', 'C': 'Only (A)-(D) and (B)-(C)', 'D': 'Only (A)-(B)', 'E': 'None of these'},
    26: {'A': 'Only (A)-(D) and (C)-(B)', 'B': 'Only (B)-(D)', 'C': 'Only (A)-(C) and (B)-(D)', 'D': 'Only (A)-(D)', 'E': 'None of these'}
}

Q39_MANUAL_OPTIONS = {
    'A': 'Only (A)',
    'B': 'Only (B)',
    'C': 'Both (A) and (B)',
    'D': 'All of these',
    'E': 'Both (A) and (C)'
}

Q45_MANUAL_OPTIONS = {
    'A': '200',
    'B': '219',
    'C': '210',
    'D': '190',
    'E': '225'
}

def clean_str(s):
    if not isinstance(s, str):
        return s
    s = s.strip()
    if "' '" in s or (s.startswith("'") and s.endswith("'") and s.count("'") > 3):
        tokens = re.findall(r"'([^']*)'", s)
        if tokens:
            return "".join(tokens).strip()
        s = s.replace("'", "")
    return s.strip()

def apply_latex_formatting(text):
    if not isinstance(text, str) or not text:
        return text
    t = text.strip()
    if '$' in t or '![' in t or '|' in t:
        return t
    
    # Quadratic equations
    t = re.sub(r'([I|V|X]+)\s*:\s*([xXyYzZ])\^2\s*([\+\-])\s*(\d+)\2\s*([\+\-])\s*(\d+)\s*=\s*0', r'\1: $\2^2 \3 \4\2 \5 \6 = 0$', t)
    t = re.sub(r'\b([xXyYzZ])\^2\s*([\+\-])\s*(\d+)\1\s*([\+\-])\s*(\d+)\s*=\s*0\b', r'$\1^2 \2 \3\1 \4 \5 = 0$', t)
    
    # Variables
    t = re.sub(r'\bRs\.?\s*([xXyYzZ])\b', r'Rs. $\1$', t)
    t = re.sub(r'\bRs\.?\s*\(([xXyYzZ]\s*[\+\-\*\/]\s*\d+)\)', r'Rs. $(\1)$', t)
    
    if re.search(r'\d+\s*[÷×]\s*\d+', t):
        t = re.sub(r'÷', r'\\div ', t)
        t = re.sub(r'×', r'\\times ', t)
        
    return t

def parse_docx_precise(docx_path):
    doc = docx.Document(docx_path)
    text_blocks = []
    
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt: text_blocks.append(txt)
        
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                txt = c.text.strip()
                if txt: text_blocks.append(txt)
                
    full_text = "\n\n".join(text_blocks)
    
    sol_split = re.split(r'\n\n(?=S1\s*\.\s*Ans|Solutions|ANSWERS)', full_text, flags=re.IGNORECASE)
    q_section = sol_split[0]
    
    q_dict = {}
    matches = list(re.finditer(r'\bQ(\d+)[\.\s]*', q_section, re.IGNORECASE))
    
    for idx, m in enumerate(matches):
        q_num = int(m.group(1))
        start_pos = m.end()
        end_pos = matches[idx+1].start() if idx + 1 < len(matches) else len(q_section)
        
        q_chunk = q_section[start_pos:end_pos].strip()
        
        opt_matches = list(re.finditer(r'\(([a-eA-E])\)\s*', q_chunk))
        if opt_matches:
            q_text = q_chunk[:opt_matches[0].start()].strip()
            options = {}
            for o_idx, om in enumerate(opt_matches):
                o_char = om.group(1).upper()
                o_start = om.end()
                o_end = opt_matches[o_idx+1].start() if o_idx + 1 < len(opt_matches) else len(q_chunk)
                o_val = q_chunk[o_start:o_end].strip()
                options[o_char] = o_val
            q_dict[q_num] = {
                "question": q_text,
                "options": options
            }
        else:
            q_dict[q_num] = {
                "question": q_chunk,
                "options": {}
            }
            
    return q_dict

def process_test(test_num):
    fname = f"sbipo_test_{test_num}.json"
    p_orig = os.path.join(orig_dir, fname)
    p_target = os.path.join(target_dir, fname)
    docx_path = os.path.join(docx_dir, docx_files[test_num - 1])
    
    with open(p_orig, 'r', encoding='utf-8') as f:
        d_orig = json.load(f)
    with open(p_target, 'r', encoding='utf-8') as f:
        d_target = json.load(f)
        
    orig_map = {q['id']: q for q in d_orig}
    target_map = {q['id']: q for q in d_target}
    docx_map = parse_docx_precise(docx_path)
    
    repaired_questions = []
    
    for qid in range(1, 101):
        q_orig = orig_map.get(qid, {})
        q_target = target_map.get(qid, {})
        q_docx = docx_map.get(qid, {})
        
        base_q = dict(q_target) if q_target else (dict(q_orig) if q_orig else {})
        base_q['id'] = qid
        
        # 1. Question Text: prefer docx text for exact alignment
        t_text = clean_str(q_target.get('question', ''))
        o_text = clean_str(q_orig.get('question', ''))
        d_text = clean_str(q_docx.get('question', ''))
        
        final_q_text = ""
        if qid == 45 and test_num in [9, 10]:
            final_q_text = "Out of the total population of all the cities 35% are females and 55% of the total population of city D are females. Total females in all the cities except D is how many more/less than total population of A."
        elif d_text and not d_text.startswith("[Question") and "text not found" not in d_text.lower():
            final_q_text = d_text
        elif t_text and not t_text.startswith("[Question") and "text not found" not in t_text.lower() and not t_text.startswith("Wiho h"):
            final_q_text = t_text
        elif o_text and not o_text.startswith("[Question") and "text not found" not in o_text.lower() and not o_text.startswith("Wiho h"):
            final_q_text = o_text
        else:
            final_q_text = f"Question {qid}"
            
        base_q['question'] = apply_latex_formatting(final_q_text)
        
        # 2. Options
        opts_t = {o['id']: clean_str(o.get('text', '')) for o in q_target.get('options', []) if 'id' in o}
        opts_o = {o['id']: clean_str(o.get('text', '')) for o in q_orig.get('options', []) if 'id' in o}
        opts_d = q_docx.get('options', {})
        
        final_options = []
        for opt_char in ['A', 'B', 'C', 'D', 'E']:
            vt = opts_t.get(opt_char, '')
            vo = opts_o.get(opt_char, '')
            vd = opts_d.get(opt_char, '')
            
            val = ""
            if test_num == 2 and qid in TEST2_MANUAL_OPTIONS:
                val = TEST2_MANUAL_OPTIONS[qid].get(opt_char, '')
            elif qid == 39 and test_num in [9, 10]:
                val = Q39_MANUAL_OPTIONS.get(opt_char, '')
            elif qid == 45 and test_num in [9, 10]:
                val = Q45_MANUAL_OPTIONS.get(opt_char, '')
            elif qid in [59, 60, 61] and (test_num in [9, 10] or "equation" in (clean_str(base_q.get('direction') or '')).lower()):
                val = DIR_OPTIONS_59_61.get(opt_char, '')
            elif qid in [76, 77, 78] and test_num in [9, 10]:
                val = DIR_OPTIONS_SYLLOGISM.get(opt_char, '')
            elif vd and vd != "" and "[Option" not in vd:
                val = vd
            elif vt and vt != "" and "[Option" not in vt:
                val = vt
            elif vo and vo != "" and "[Option" not in vo:
                val = vo
            else:
                val = f"[Option {opt_char}]"
                
            val = apply_latex_formatting(val)
            
            # preserve option image
            opt_img = None
            for opt_obj in q_target.get('options', []):
                if opt_obj.get('id') == opt_char and opt_obj.get('image'):
                    opt_img = opt_obj.get('image')
                    break
            if not opt_img:
                for opt_obj in q_orig.get('options', []):
                    if opt_obj.get('id') == opt_char and opt_obj.get('image'):
                        opt_img = opt_obj.get('image')
                        break
                        
            final_options.append({
                "id": opt_char,
                "text": val,
                "image": opt_img
            })
            
        base_q['options'] = final_options
        
        # 3. correctAnswer
        ca_t = q_target.get('correctAnswer')
        ca_o = q_orig.get('correctAnswer')
        base_q['correctAnswer'] = ca_t if ca_t in ['A','B','C','D','E'] else (ca_o if ca_o in ['A','B','C','D','E'] else "A")
        
        # 4. Explanation
        exp_t = clean_str(q_target.get('explanation'))
        exp_o = clean_str(q_orig.get('explanation'))
        base_q['explanation'] = apply_latex_formatting(exp_t if exp_t else (exp_o if exp_o else ""))
        
        # 5. Direction
        dir_t = clean_str(q_target.get('direction'))
        dir_o = clean_str(q_orig.get('direction'))
        base_q['direction'] = apply_latex_formatting(dir_t if dir_t else dir_o)
        
        # Metadata
        base_q['exam'] = clean_str(q_target.get('exam')) or clean_str(q_orig.get('exam')) or "SBI PO Prelims"
        base_q['year'] = q_target.get('year') or q_orig.get('year') or 2024
        base_q['subject'] = clean_str(q_target.get('subject')) or clean_str(q_orig.get('subject')) or ""
        base_q['topic'] = clean_str(q_target.get('topic')) or clean_str(q_orig.get('topic')) or ""
        base_q['difficulty'] = clean_str(q_target.get('difficulty')) or clean_str(q_orig.get('difficulty')) or "Medium"
        base_q['marks'] = q_target.get('marks') or q_orig.get('marks') or 1
        base_q['negativeMarks'] = q_target.get('negativeMarks') or q_orig.get('negativeMarks') or 0.25
        base_q['questionImage'] = q_target.get('questionImage') or q_orig.get('questionImage')
        
        repaired_questions.append(base_q)
        
    # Write repaired output to BOTH target_dir and orig_dir
    with open(p_target, 'w', encoding='utf-8') as f:
        json.dump(repaired_questions, f, indent=2, ensure_ascii=False)
    with open(p_orig, 'w', encoding='utf-8') as f:
        json.dump(repaired_questions, f, indent=2, ensure_ascii=False)
        
    print(f"Successfully fixed and saved {fname} (100 questions).")

for test_num in range(1, 11):
    process_test(test_num)

print("\nMaster fix completed for all 10 SBI PO tests!")

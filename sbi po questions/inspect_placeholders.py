import docx
import re

specs = {
    2: [24, 25, 26],
    9: [39, 59, 60, 61],
    10: [39, 45, 59, 60, 61, 76, 77, 78]
}

docx_files = {
    2: "SBI-PO-Pre-2022-20th-Dec-Shift-Wise-Previous-Year-Paper-Mock-6.docx",
    9: "SBI-PO-Pre-2025-Memory-Based-Paper-Based-on-4th-August-1st-Shift (1).docx",
    10: "SBI-PO-Pre-2025-Memory-Based-Paper-Based-on-4th-August-1st-Shift.docx"
}

docx_dir = r"C:\Users\Administrator\Downloads\sbi po prelims"

for t_num, q_list in specs.items():
    d_file = docx_files[t_num]
    d_path = f"{docx_dir}\\{d_file}"
    doc = docx.Document(d_path)
    print(f"\n================ Test {t_num} ({d_file}) ================")
    
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    for qid in q_list:
        print(f"\n--- Searching Q{qid} ---")
        pat = rf'^\s*Q\s*\.?\s*{qid}\b'
        found = False
        for idx, p in enumerate(paragraphs):
            if re.search(pat, p, re.IGNORECASE):
                print("\n".join(paragraphs[idx:min(len(paragraphs), idx+8)]))
                found = True
                break
        if not found:
            for t_idx, t in enumerate(doc.tables):
                for r_idx, r in enumerate(t.rows):
                    for c_idx, c in enumerate(r.cells):
                        if re.search(pat, c.text, re.IGNORECASE):
                            print(f"Table {t_idx} R{r_idx} C{c_idx}:")
                            print(c.text[:400])
                            found = True
                            break
                    if found: break
                if found: break


import json
import os
import re
import docx

docx_dir = r"C:\Users\Administrator\Downloads\sbi po prelims"
docx_files = [
    "SBI-PO-Pre-2022-19th-Dec-Shift-Wise-Previous-Year-Paper-Mock-5.docx",
    "SBI-PO-Pre-2022-20th-Dec-Shift-Wise-Previous-Year-Paper-Mock-6.docx",
    "SBI-PO-Pre-2024-25-Memory-Based-Paper-16-Mar-2025-1st-shift.docx",
    "SBI-PO-Pre-2024-25-Memory-Based-Paper-24-Mar-2025-1st-shift-1.docx",
    "SBI-PO-Pre-2024-25-Memory-Based-Paper-8-Mar-2025-1st-shift.docx",
    "SBI-PO-Pre-2024-25-Memory-Based-Paper-8-March-2025-3rd-shift.docx",
    "SBI-PO-Pre-2024-25-Memory-Based-Paper-8-March-2025-4th-shift.docx",
    "SBI-PO-Pre-2024-25-Memory-Based-Question-Paper-8-Mar-2025-2nd-shift-1.docx",
    "SBI-PO-Pre-2025-Memory-Based-Paper-Based-on-4th-August-1st-Shift (1).docx",
    "SBI-PO-Pre-2025-Memory-Based-Paper-Based-on-4th-August-1st-Shift.docx"
]

def parse_docx_precise(docx_path):
    doc = docx.Document(docx_path)
    text_blocks = []
    
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt: text_blocks.append(txt)
        
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                txt = c.text.strip()
                if txt: text_blocks.append(txt)
                
    full_text = "\n\n".join(text_blocks)
    
    # Split solutions out if present
    sol_split = re.split(r'\n\n(?=S1\s*\.\s*Ans|Solutions|ANSWERS)', full_text, flags=re.IGNORECASE)
    q_section = sol_split[0]
    
    q_dict = {}
    matches = list(re.finditer(r'\bQ(\d+)[\.\s]*', q_section, re.IGNORECASE))
    
    for idx, m in enumerate(matches):
        q_num = int(m.group(1))
        start_pos = m.end()
        end_pos = matches[idx+1].start() if idx + 1 < len(matches) else len(q_section)
        
        q_chunk = q_section[start_pos:end_pos].strip()
        
        opt_matches = list(re.finditer(r'\(([a-eA-E])\)\s*', q_chunk))
        if opt_matches:
            q_text = q_chunk[:opt_matches[0].start()].strip()
            options = {}
            for o_idx, om in enumerate(opt_matches):
                o_char = om.group(1).upper()
                o_start = om.end()
                o_end = opt_matches[o_idx+1].start() if o_idx + 1 < len(opt_matches) else len(q_chunk)
                o_val = q_chunk[o_start:o_end].strip()
                options[o_char] = o_val
            q_dict[q_num] = {
                "question": q_text,
                "options": options
            }
        else:
            q_dict[q_num] = {
                "question": q_chunk,
                "options": {}
            }
            
    return q_dict

for test_num in range(1, 11):
    docx_path = os.path.join(docx_dir, docx_files[test_num - 1])
    parsed = parse_docx_precise(docx_path)
    print(f"Test {test_num}: parsed {len(parsed)} questions from DOCX (Qs: {min(parsed.keys()) if parsed else 0} to {max(parsed.keys()) if parsed else 0}).")
